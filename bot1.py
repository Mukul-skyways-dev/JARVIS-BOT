# ============================================================
#  JARVIS - A AERO CROWN OFFICIAL BOT  |  V3.0.1
#  Single-file edition — all imports, all features inline
# ============================================================

# ── Standard Library ─────────────────────────────────────────
import os
import io
import csv
import time
import json
import asyncio
import random
import textwrap
from datetime import datetime, timedelta
from threading import Thread

# ── Third-party ───────────────────────────────────────────────
import discord
from discord.ext import commands, tasks
from discord.ui import Modal, TextInput, View, Button

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

import sqlite3
import requests
import pytz
from openai import OpenAI
from flask import Flask

# ── PDF / DOCX (optional, graceful fallback) ─────────────────
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as rlcanvas
    from reportlab.lib import colors as rl_colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ─────────────────────────────────────────────────────────────
#  KEEP-ALIVE (Flask)
# ─────────────────────────────────────────────────────────────
app = Flask('')

@app.route('/')
def home():
    return "JARVIS is alive ✈"

def _run_flask():
    app.run(host='0.0.0.0', port=8080)

def keep_alive():
    t = Thread(target=_run_flask, daemon=True)
    t.start()

# ─────────────────────────────────────────────────────────────
#  CONFIG
# ─────────────────────────────────────────────────────────────
TOKEN          = os.getenv("TOKEN")
GROQ_KEY       = os.getenv("GROQ_API_KEY")
CHANNEL_ID     = int(os.getenv("CHANNEL_ID", "0"))
FUEL_CH_ID     = int(os.getenv("FUEL_CHANNEL_ID", str(CHANNEL_ID)))
SHARE_CH_ID    = int(os.getenv("SHARE_CHANNEL_ID", str(CHANNEL_ID)))
ADMIN_ROLE_NAME = "Admin"
BOT_VERSION    = "V3.0.1 ALPHA"
FOOTER         = f"JARVIS - A AERO CROWN OFFICIAL BOT • {BOT_VERSION}"

IST = pytz.timezone("Asia/Kolkata")

groq = OpenAI(api_key=GROQ_KEY, base_url="https://api.groq.com/openai/v1")

intents = discord.Intents.default()
intents.message_content = True
intents.members = True
bot = commands.Bot(command_prefix="!", intents=intents)

# ─────────────────────────────────────────────────────────────
#  DATABASE — AM4 (read-only analytics)
# ─────────────────────────────────────────────────────────────
DB_URL  = "https://github.com/Mukul-skyways-dev/JARVIS-BOT/releases/download/Dv1/am4_data.db.updated"
DB_FILE = "am4_data.db"

def download_db():
    print("⬇ Downloading AM4 database …")
    try:
        r = requests.get(DB_URL, timeout=30)
        r.raise_for_status()
        with open(DB_FILE, "wb") as f:
            f.write(r.content)
        print("✅ Database ready")
    except Exception as e:
        print("❌ DB download failed:", e)

download_db()

conn   = sqlite3.connect(DB_FILE, check_same_thread=False)
conn.row_factory = sqlite3.Row
cursor = conn.cursor()

# ─────────────────────────────────────────────────────────────
#  DATABASE — Dynamic (all new features)
# ─────────────────────────────────────────────────────────────
conn_dyn = sqlite3.connect("new_am4.db", check_same_thread=False)
conn_dyn.row_factory = sqlite3.Row
cur = conn_dyn.cursor()   # alias

def _dyn(sql, params=()):
    cur.execute(sql, params)
    conn_dyn.commit()

# ── Schema ────────────────────────────────────────────────────
cur.executescript("""
CREATE TABLE IF NOT EXISTS users (
    user_id TEXT PRIMARY KEY, username TEXT
);
CREATE TABLE IF NOT EXISTS shares (
    user_id TEXT, value REAL, date TEXT, window_id TEXT
);
CREATE TABLE IF NOT EXISTS activity (
    user_id TEXT PRIMARY KEY,
    total INT DEFAULT 0, attended INT DEFAULT 0, missed INT DEFAULT 0,
    streak INT DEFAULT 0, last_window TEXT, miss_streak INT DEFAULT 0
);
CREATE TABLE IF NOT EXISTS fuel_data (
    day TEXT, time TEXT, fuel REAL, co2 REAL
);
CREATE TABLE IF NOT EXISTS fuel_schedules (
    user_id TEXT PRIMARY KEY, threshold REAL
);
CREATE TABLE IF NOT EXISTS alliance_members (
    user_id TEXT PRIMARY KEY, username TEXT, airline TEXT,
    rank TEXT, pax REAL, revenue REAL, joined TEXT
);
CREATE TABLE IF NOT EXISTS votes (
    poll_id TEXT, user_id TEXT, choice TEXT,
    PRIMARY KEY (poll_id, user_id)
);
CREATE TABLE IF NOT EXISTS polls (
    poll_id TEXT PRIMARY KEY, question TEXT, options TEXT,
    created TEXT, ends TEXT, channel_id INTEGER, msg_id INTEGER,
    image_url TEXT, closed INTEGER DEFAULT 0
);
CREATE TABLE IF NOT EXISTS ai_memory (
    user_id TEXT, role TEXT, content TEXT, ts TEXT
);
CREATE TABLE IF NOT EXISTS bot_usage (
    user_id TEXT PRIMARY KEY, username TEXT,
    points INTEGER DEFAULT 0, last_used REAL DEFAULT 0
);
""")
conn_dyn.commit()

# ─────────────────────────────────────────────────────────────
#  UTILS
# ─────────────────────────────────────────────────────────────
def clean(x):   return str(x).replace(",","").replace('"',"").replace("'","").strip()
def to_int(x):
    try: return int(float(clean(x)))
    except: return 0
def to_float(x):
    try: return float(clean(x))
    except: return 0.0
def norm(x):    return x.upper().replace("-","").replace(" ","")
def money(x):   return f"${x:,.0f}"
def now_ist():  return datetime.now(IST)
def ts():       return now_ist().strftime("%d %b %Y • %I:%M %p IST")
def window_id():return now_ist().date().isoformat()

def parse_money(v):
    try:
        v = v.lower().replace(",","").strip()
        if "b" in v: return float(v.replace("b",""))*1e9
        if "m" in v: return float(v.replace("m",""))*1e6
        return float(v)
    except: return None

def format_time(hours):
    h = int(hours); m = int((hours-h)*60)
    s = int((((hours-h)*60)-m)*60)
    return f"{h:02}:{m:02}:{s:02} ({round(hours,3)} hr)"

def is_admin(member):
    if not member: return False
    return member.guild_permissions.manage_guild

# ─────────────────────────────────────────────────────────────
#  GRAY GRAPH HELPER  (all graphs: gray bg)
# ─────────────────────────────────────────────────────────────
GRAPH_BG  = "#2f2f2f"
GRAPH_FIG = "#1e1e1e"
GRAPH_GRID= "#444444"
GRAPH_TEXT= "#e0e0e0"

def gray_fig(w=9, h=5):
    fig, ax = plt.subplots(figsize=(w, h))
    fig.patch.set_facecolor(GRAPH_FIG)
    ax.set_facecolor(GRAPH_BG)
    ax.tick_params(colors=GRAPH_TEXT)
    ax.xaxis.label.set_color(GRAPH_TEXT)
    ax.yaxis.label.set_color(GRAPH_TEXT)
    ax.title.set_color(GRAPH_TEXT)
    for sp in ax.spines.values():
        sp.set_color(GRAPH_GRID)
    ax.grid(alpha=0.2, linestyle=':', color=GRAPH_GRID)
    return fig, ax

def save_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=150,
                facecolor=fig.get_facecolor())
    buf.seek(0)
    plt.close(fig)
    return buf

# ─────────────────────────────────────────────────────────────
#  EXPORT VIEW  (inline — CSV + DOCX + PDF)
# ─────────────────────────────────────────────────────────────
class ExportView(View):
    def __init__(self, data: dict, title: str = "JARVIS Report"):
        super().__init__(timeout=300)
        self.data  = data
        self.title = title

    # ── CSV ───────────────────────────────────────────────────
    @discord.ui.button(label="📄 CSV", style=discord.ButtonStyle.secondary)
    async def export_csv(self, interaction: discord.Interaction, button: Button):
        buf = io.StringIO()
        w   = csv.writer(buf)
        w.writerow(["Field", "Value"])
        for k, v in self.data.items():
            w.writerow([k, v])
        buf.seek(0)
        await interaction.response.send_message(
            f"📄 **{self.title}** — CSV Export\n`{ts()}`",
            file=discord.File(io.BytesIO(buf.getvalue().encode()), "report.csv"),
            ephemeral=True
        )

    # ── DOCX ──────────────────────────────────────────────────
    @discord.ui.button(label="📝 DOCX", style=discord.ButtonStyle.blurple)
    async def export_docx(self, interaction: discord.Interaction, button: Button):
        if not HAS_DOCX:
            return await interaction.response.send_message(
                "❌ python-docx not installed", ephemeral=True)
        doc = DocxDoc()
        doc.add_heading(self.title, 0)
        doc.add_paragraph(f"Generated: {ts()}")
        doc.add_paragraph("")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = "Field"; hdr[1].text = "Value"
        for k, v in self.data.items():
            row = tbl.add_row().cells
            row[0].text = str(k); row[1].text = str(v)
        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        await interaction.response.send_message(
            f"📝 **{self.title}** — DOCX Export\n`{ts()}`",
            file=discord.File(buf, "report.docx"), ephemeral=True)

    # ── PDF ───────────────────────────────────────────────────
    @discord.ui.button(label="📕 PDF", style=discord.ButtonStyle.red)
    async def export_pdf(self, interaction: discord.Interaction, button: Button):
        if not HAS_PDF:
            return await interaction.response.send_message(
                "❌ reportlab not installed", ephemeral=True)
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter)
        styles = getSampleStyleSheet()
        elems  = [Paragraph(self.title, styles['Title']),
                  Paragraph(f"Generated: {ts()}", styles['Normal']),
                  Spacer(1, 12)]
        rows = [["Field", "Value"]] + [[str(k), str(v)] for k,v in self.data.items()]
        t = Table(rows, colWidths=[200, 300])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0),(-1,0), rl_colors.grey),
            ('TEXTCOLOR',  (0,0),(-1,0), rl_colors.whitesmoke),
            ('GRID',       (0,0),(-1,-1), 0.5, rl_colors.black),
            ('FONTSIZE',   (0,0),(-1,-1), 9),
        ]))
        elems.append(t)
        doc.build(elems)
        buf.seek(0)
        await interaction.response.send_message(
            f"📕 **{self.title}** — PDF Export\n`{ts()}`",
            file=discord.File(buf, "report.pdf"), ephemeral=True)

# ─────────────────────────────────────────────────────────────
#  DIFFICULTY SYSTEM
# ─────────────────────────────────────────────────────────────
def get_user_mode(user_id):
    cursor.execute(
        "SELECT difficulty FROM player_settings WHERE user_id=?", (str(user_id),))
    row = cursor.fetchone()
    return row[0].lower() if row and row[0] else "realism"

def set_user_mode(user_id, mode):
    cursor.execute(
        "INSERT OR REPLACE INTO player_settings (user_id, difficulty) VALUES (?,?)",
        (str(user_id), mode))
    conn.commit()

# ─────────────────────────────────────────────────────────────
#  FETCH ROUTE / PLANE
# ─────────────────────────────────────────────────────────────
def get_route(frm, to):
    cursor.execute("""
        SELECT * FROM routes
        WHERE (f_iata=? AND t_iata=?) OR (f_iata=? AND t_iata=?)
        LIMIT 1
    """, (frm.upper(), to.upper(), to.upper(), frm.upper()))
    row = cursor.fetchone()
    if not row: return None
    return {"distance": to_float(row[5]), "y": to_int(row[9]),
            "j": to_int(row[10]), "f": to_int(row[11]), "cargo": to_int(row[8])}

def get_all_planes():
    cursor.execute(
        "SELECT model, variant, capacity, range, speed, fuel_efficiency, cost FROM aircraft")
    return [{"name": f"{r[0]} {r[1]}", "capacity": to_int(r[2]),
             "range": to_float(r[3]), "speed": to_float(r[4]),
             "fuel": to_float(r[5]), "cost": to_int(r[6])}
            for r in cursor.fetchall()]

def get_plane(name):
    key = norm(name)
    for p in get_all_planes():
        if key in norm(p["name"]):
            return p
    return None

def airport_name(iata):
    iata = iata.upper()
    for q in [
        "SELECT city, country FROM routes WHERE iata=? LIMIT 1",
        "SELECT city, country FROM routes WHERE f_iata=? OR t_iata=? LIMIT 1"
    ]:
        try:
            args = (iata,) if q.count("?") == 1 else (iata, iata)
            cursor.execute(q, args)
            row = cursor.fetchone()
            if row and row[0] and row[1]:
                return f"{iata} — {row[0]}, {row[1]}"
        except: pass
    return iata

# ─────────────────────────────────────────────────────────────
#  CALC ENGINE V3
# ─────────────────────────────────────────────────────────────
def calc(route, plane, user_id, mods=None):
    mode  = get_user_mode(user_id)
    dist  = float(route["distance"])
    speed = float(plane["speed"])
    if mods and "speed" in mods: speed *= 1.1
    time_h = dist / speed if speed else 1
    trips  = max(1, int(24 / time_h))

    y, j, f = int(route["y"]), int(route["j"]), int(route["f"])
    total   = y + j + f
    cap     = int(plane["capacity"])

    if mode == "easy":
        lf = 1.0
        y_price = 0.4*dist + 170; j_price = 0.8*dist + 560; f_price = 1.2*dist + 1200
        fuel_mult=4; co2_mult=1.8; acheck=20000; repair=15000; cargo_mul=0.5
    else:
        lf = 0.85
        y_price = 0.3*dist + 150; j_price = 0.6*dist + 500; f_price = 0.9*dist + 1000
        fuel_mult=5.5; co2_mult=2.5; acheck=40000; repair=25000; cargo_mul=0.35

    if total > 0:
        y_c = int(cap*(y/total)*lf); j_c = int(cap*(j/total)*lf)
        f_c = max(0, cap - y_c - j_c)
    else:
        y_c = j_c = f_c = 0

    income_trip  = y_c*y_price + j_c*j_price + f_c*f_price
    cargo_income = float(route.get("cargo",0)) * cargo_mul
    income_trip += cargo_income

    fuel = dist * float(plane["fuel"]) * fuel_mult
    co2  = dist * co2_mult
    if mods:
        if "fuel" in mods: fuel *= 0.9
        if "co2"  in mods: co2  *= 0.9

    total_cost  = fuel + co2 + acheck + repair
    profit_trip = income_trip - total_cost
    ci          = int((profit_trip/income_trip)*100) if income_trip else 0

    return {
        "mode": mode, "distance": int(dist), "time": round(time_h,2), "trips": trips,
        "y": y_c, "j": j_c, "f": f_c,
        "y_price": int(y_price), "j_price": int(j_price), "f_price": int(f_price),
        "income_trip": int(income_trip), "cargo_income": int(cargo_income),
        "fuel": int(fuel), "fuel_lb": int(fuel*2.2),
        "co2":  int(co2),  "co2_q":   int(co2*1.1),
        "acheck": int(acheck), "repair": int(repair), "total_cost": int(total_cost),
        "profit_trip": int(profit_trip), "ci": ci,
        "income_day":  int(income_trip*trips), "fuel_day":   int(fuel*trips),
        "co2_day":     int(co2*trips),          "profit_day": int(profit_trip*trips)
    }

# ─────────────────────────────────────────────────────────────
#  USAGE TRACKER
# ─────────────────────────────────────────────────────────────
COOLDOWN = 3
def add_usage(user):
    now = time.time()
    cur.execute("SELECT last_used FROM bot_usage WHERE user_id=?", (str(user.id),))
    row = cur.fetchone()
    if row and now - row[0] < COOLDOWN: return
    cur.execute("""
        INSERT INTO bot_usage (user_id,username,points,last_used) VALUES (?,?,1,?)
        ON CONFLICT(user_id) DO UPDATE SET
        points=points+1, username=excluded.username, last_used=excluded.last_used
    """, (str(user.id), user.name, now))
    conn_dyn.commit()

@bot.event
async def on_command(ctx):
    add_usage(ctx.author)

# ─────────────────────────────────────────────────────────────
#  SHARE TRACKING HELPERS
# ─────────────────────────────────────────────────────────────
async def register_user(user):
    uid = str(user.id)
    cur.execute("SELECT * FROM users WHERE user_id=?", (uid,))
    if not cur.fetchone():
        cur.execute("INSERT INTO users VALUES (?,?)", (uid, str(user)))
        cur.execute("INSERT INTO activity VALUES (?,0,0,0,0,NULL,0)", (uid,))
        conn_dyn.commit()

def rank_label(cons):
    if cons >= 90: return "🏆 Legend"
    if cons >= 75: return "💎 Diamond"
    if cons >= 60: return "🥇 Gold"
    if cons >= 40: return "🥈 Silver"
    return "🥉 Bronze"

# ─────────────────────────────────────────────────────────────
#  SHARE TRACKING STATE
# ─────────────────────────────────────────────────────────────
current_window = {"id": None, "open_time": None}

# ─────────────────────────────────────────────────────────────
#  MENU
# ─────────────────────────────────────────────────────────────
class EliteMenu(View):
    def __init__(self):
        super().__init__(timeout=None)

    def _emb(self, title, desc, color):
        e = discord.Embed(title=title, description=desc, color=color)
        e.set_footer(text=f"{FOOTER} • {ts()}")
        return e

    @discord.ui.button(label="✈ Route System", style=discord.ButtonStyle.blurple)
    async def route_help(self, i, b):
        await i.response.send_message(embed=self._emb("✈ Route Command",
            "`!route DEL BOM A320`\n\n📊 Includes:\n• Flight time, distance, CI\n"
            "• Demand (Y/J/F) • Config & Ticket\n• A-check, Repair\n"
            "• Profit (Trip + Day) • Mods & Stopover", 0x3498db), ephemeral=True)

    @discord.ui.button(label="🔥 Best Routes", style=discord.ButtonStyle.red)
    async def best_help(self, i, b):
        await i.response.send_message(embed=self._emb("🔥 Best Route Finder",
            "`!best_r DEL A320`\n`!best_short DEL A320`\n`!best_long DEL A320`\n\n"
            "📈 Shows: Profit, CI, Trips, Demand, Flight Time", 0xe74c3c), ephemeral=True)

    @discord.ui.button(label="⚖ Compare Planes", style=discord.ButtonStyle.gray)
    async def compare_help(self, i, b):
        await i.response.send_message(embed=self._emb("⚖ Plane Comparison",
            "`!compare A320 vs B737`\n\n📊 Shows:\n• Cost, Capacity, Range, Speed\n"
            "• Fuel, CO2, Income, Profit, CI\n• Radar + Performance Graph", 0x95a5a6),
            ephemeral=True)

    @discord.ui.button(label="⛽ Fuel", style=discord.ButtonStyle.secondary)
    async def fuel_help(self, i, b):
        await i.response.send_message(embed=self._emb("⛽ Fuel Commands",
            "`!fuel` — Current prediction (DB-based)\n`!fuelgraph` — Graph\n"
            "`!fuelschedule <amount>` — Personal DM alert\n\n"
            "Auto alert every 30min (AM4 cycle)", 0xf39c12), ephemeral=True)

    @discord.ui.button(label="🏦 Alliance", style=discord.ButtonStyle.green)
    async def alliance_help(self, i, b):
        await i.response.send_message(embed=self._emb("🏦 Alliance Commands",
            "`!allianceadd` — Register (modal form)\n`!alliancestats [@user]` — Stats\n"
            "`!alliancecompare @u1 @u2` — Head-to-head\n`!allianceboard` — Leaderboard\n"
            "`!alliancegraph` — Graph\n`!alliancehistory` — History\n"
            "`!allianceexport` — Export CSV/DOCX/PDF\n\n⚠ Admin-only management", 0x2ecc71),
            ephemeral=True)

    @discord.ui.button(label="🗳 Voting", style=discord.ButtonStyle.blurple)
    async def vote_help(self, i, b):
        await i.response.send_message(embed=self._emb("🗳 Voting System",
            "`!createpoll` — Create poll (modal, 5hr, hidden votes)\n"
            "`!checkvote <poll_id>` — Check your vote\n"
            "`!pollexport <poll_id>` — Export results\n\n"
            "• Voter ID DM'd on vote\n• Auto-results after 5hr", 0x9b59b6), ephemeral=True)

    @discord.ui.button(label="🤖 AI / General", style=discord.ButtonStyle.secondary)
    async def general_help(self, i, b):
        await i.response.send_message(embed=self._emb("🤖 General Commands",
            "`!ask <question>` — AI with memory\n`!clearhistory` — Clear AI memory\n"
            "`!menu` `!ping` `!difficulty easy/realism`\n"
            "`!announce <msg>` — Admin DM blast\n"
            "`!leaderboard` — Usage board\n\n"
            "Chat: Hi / Hello / Jarvis (mention or DM)", 0
