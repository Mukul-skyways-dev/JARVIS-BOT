import csv
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# =========================
# CSV EXPORT
# =========================
def export_csv(report_data, filename="report.csv"):

    with open(filename, "w", newline="", encoding="utf-8") as f:

        writer = csv.writer(f)

        writer.writerow(["Field", "Value"])

        for k, v in report_data.items():
            writer.writerow([k, v])

    return filename

# =========================
# DOCX EXPORT
# =========================
def export_docx(report_data, filename="report.docx"):

    doc = Document()

    doc.add_heading("JARVIS Report", level=1)

    for k, v in report_data.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.save(filename)

    return filename

# =========================
# PDF EXPORT
# =========================
def export_pdf(report_data, filename="report.pdf"):

    doc = SimpleDocTemplate(filename)

    styles = getSampleStyleSheet()

    elements = []

    elements.append(
        Paragraph("JARVIS Analytics Report", styles['Title'])
    )

    elements.append(Spacer(1, 12))

    for k, v in report_data.items():

        text = f"<b>{k}</b>: {v}"

        elements.append(
            Paragraph(text, styles['BodyText'])
        )

        elements.append(Spacer(1, 6))

    doc.build(elements)

    return filename
